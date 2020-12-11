from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Cm

document = Document()
sections = document.sections
for section in sections:
  section.top_margin = Cm(3)
  section.right_margin = Cm(2)
  section.bottom_margin = Cm(2)
  section.left_margin = Cm(3)

print('Nenhum dos dados colocados aqui são salvos externamente ou enviados para o desenvolvedor, tudo é completamente anônimo e só fica salvo no seu computador.')

file = input('Insira o nome que quer no arquivo: ')
school = input('Insira o nome do seu colégio: ')
name = input('Insira seu nome: ')
city = input('Insira sua cidade: ')
year = input('Insira o ano de realização do trabalho: ')
text = input('Insira o texto do seu trabalho: ')

def space():
  spacing = paragraph.paragraph_format
  spacing.space_before = Pt(0)
  spacing.space_after = Pt(0)

paragraph.style = document.styles.add_style('ABNT', WD_STYLE_TYPE.PARAGRAPH)
paragraph = document.add_paragraph(school)
space()
paragraph = document.add_paragraph(name)
space()
paragraph = document.add_paragraph(city)
space()
paragraph = document.add_paragraph(year)
space()
paragraph = document.add_paragraph(text)
space()

font = paragraph.style.font
font.name = 'Arial'
font.size = Pt(12)

document.save(file + ".docx")